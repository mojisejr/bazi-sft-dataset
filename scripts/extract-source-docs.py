#!/usr/bin/env python3
"""Extract structured Bazi source data from docs/ originals into knownlage/extracted/*.txt.

Dev tool (run manually): regenerates the deterministic knowledge files consumed by
src/lib/bazi/topic-knowledge.ts. Requires python-docx and openpyxl.

Outputs:
  - knownlage/extracted/love-day-pillar.txt   (from ความรัก xlsx, sheet "หลักวันเท่านั้น")
  - knownlage/extracted/source7-custom.txt     (deity tables 6/7 from Source7 docx §5)
  - knownlage/extracted/kheangkhung-reference.txt (ตำราเคี้ยงคุง full text, fallback ref)
"""
import os
import re
import unicodedata

import docx
import openpyxl

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS = os.path.join(ROOT, "docs")
OUT = os.path.join(ROOT, "knownlage", "extracted")

STEMS = ["甲", "乙", "丙", "丁", "戊", "己", "庚", "辛", "壬", "癸"]
BRANCHES = ["子", "丑", "寅", "卯", "辰", "巳", "午", "未", "申", "酉", "戌", "亥"]


def norm(text):
    """Normalize source encoding quirks before matching.

    เอกสารต้นฉบับเก็บกิ่ง 辰 เป็น CJK Compatibility Ideograph (U+F971) และสะกด
    "น้ำ" เป็น "น้ํา" (นิคหิต U+0E4D + สระอา U+0E32) — ถ้าไม่ normalize จะ match
    ไม่ติดทำให้ record (โดยเฉพาะหลักวันกิ่ง辰) หายไป. ดู knowledge-audit §A8.
    """
    if text is None:
        return ""
    return unicodedata.normalize("NFKC", str(text)).replace("ํา", "ำ")


def write(name, lines):
    path = os.path.join(OUT, name)
    with open(path, "w", encoding="utf-8") as fh:
        fh.write("\n".join(lines) + "\n")
    print(f"wrote {name}: {len(lines)} lines")


def extract_love_day_pillar():
    """Pair each day-stem verdict row (A) with its branch/12-qi/spouse row (B).

    Day pillars obey stem-branch parity (yang stem -> yang branch), so the 6 rows
    listed per stem are exactly the valid (dayStem, dayBranch) day pillars.
    """
    wb = openpyxl.load_workbook(
        os.path.join(DOCS, "ความรัก (หลักวัน ราศีบน-ล่าง).xlsx"),
        read_only=True, data_only=True,
    )
    ws = wb["หลักวันเท่านั้น"]
    rows = []
    for row in ws.iter_rows(values_only=True):
        cells = [norm(c).strip() for c in row if c is not None and norm(c).strip()]
        if cells:
            rows.append(cells)

    out = [
        "# love day-pillar verdicts (source: ความรัก หลักวันเท่านั้น)",
        "# format: STEM BRANCH QI | spouseText | reactionText",
    ]
    pending_stem = None
    pending_reaction = None
    for cells in rows:
        # A-row: <elem-code> <stem> <num> [reactionText]  (stem is cells[1])
        # reactionText อาจว่าง (เช่น 壬 11.0) → ต้องยอม len>=3 ไม่งั้น B-row ถัดไป (壬午) หลุดคู่
        if len(cells) >= 3 and cells[1] in STEMS and re.match(r"^\d+(\.\d+)?$", cells[2]):
            pending_stem = cells[1]
            pending_reaction = cells[3] if len(cells) >= 4 else None
            continue
        # B-row: <num> <branch> <qi> <spouseText>
        if len(cells) >= 4 and cells[1] in BRANCHES and pending_stem:
            branch, qi, spouse = cells[1], cells[2], cells[3]
            out.append(f"{pending_stem} {branch} {qi} | {spouse} | {pending_reaction or '-'}")
            pending_stem = None
            pending_reaction = None
    return out


def extract_source7_custom():
    d = docx.Document(os.path.join(DOCS, "Source7_ การเสริมดวง.docx"))
    out = [
        "# Source7 §5 custom guardian deities by chart character",
        "# format: char | deity | degree (องศา; - = ไม่ระบุ)",
    ]

    # Table 6: ราศีบน (heavenly stem) -> รายชื่อ | องศา ; Table 7: ราศีล่าง (branch) -> รายชื่อ | องศา
    def emit(table, header, keyset):
        out.append(header)
        for r in table.rows[1:]:
            cells = [norm(c.text).strip() for c in r.cells]
            if len(cells) >= 2 and cells[0] in keyset and cells[1]:
                degree = cells[2] if len(cells) >= 3 and cells[2] else "-"
                out.append(f"{cells[0]} | {cells[1]} | {degree}")

    emit(d.tables[6], "# DEITY_UPPER", set(STEMS))
    emit(d.tables[7], "# DEITY_LOWER", set(BRANCHES))
    return out


def extract_kheangkhung():
    d = docx.Document(os.path.join(DOCS, "ตำราโหราศาสตร์เคี้ยงคุง.docx"))
    out = ["# ตำราโหราศาสตร์เคี้ยงคุง (foundational reference, fallback)"]
    for p in d.paragraphs:
        t = norm(p.text).strip()
        if t:
            out.append(t)
    return out


def _find_docs(basename):
    for root, _dirs, files in os.walk(DOCS):
        for f in files:
            if f == basename and not f.startswith("~$"):
                return os.path.join(root, f)
    return None


def extract_career_relations():
    """ตาราง 12 เชี่ยงแซ → คำทำนาย ลูกน้อง/หุ้นส่วน/เจ้านาย (คู่สมพงษ์การงาน xlsx).

    แต่ละ sheet: คอลัมน์ 3 = ชื่อเชี่ยงแซ, คอลัมน์ 5 = คำทำนาย (12 แถว A1..A12).
    """
    path = _find_docs("คู่สมพงษ์(การงาน).xlsx")
    out = ["# career-relation 12-qi verdicts (source: คู่สมพงษ์การงาน)", "# format: [relation] qi | verdict"]
    if not path:
        return out
    twelve_qi = {"เชี่ยงแซ", "หมกยก", "กวงตั่ว", "ลิ่มกัว", "ตี้อ๋วง", "ซวย",
                 "แป่", "ซี่", "หมอ", "เจ๊าะ", "ทอ", "เอี้ยง"}
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    sheets = {
        "employee": "คำทำนาย ลูกน้อง>ตัวเรา",
        "partner": "คำทำนายหุ้นส่วนเพื่อนร่วมงาน",
        "boss": "คำทำนายตัวเรา>เจ้านาย",
    }
    def emit_qi_table(workbook, sheet, relation):
        """หา qi + คำทำนายแบบ dynamic (คอลัมน์ของแต่ละ sheet ไม่ตรงกัน บางอันมี leading empty col)."""
        if sheet not in workbook.sheetnames:
            return
        ws = workbook[sheet]
        for row in ws.iter_rows(min_row=1, max_row=20, values_only=True):
            vals = [norm(c).strip() for c in row if c is not None and norm(c).strip()]
            qi = next((v for v in vals if v in twelve_qi), "")
            if not qi:
                continue
            cands = [
                v for v in vals
                if v != qi and len(v) > 10
                and not re.match(r"^[\d.]+$", v)
                and not re.match(r"^A\d+$", v)
                and not v.startswith("UPDATE")
            ]
            if cands:
                out.append(f"[{relation}] {qi} | {max(cands, key=len)}")

    for relation, sheet in sheets.items():
        emit_qi_table(wb, sheet, relation)

    # คู่รัก: คู่สมพงษ์(ความรัก).xlsx sheet "12เชี่ยงแซความรัก" → คำทำนายคู่รักตาม 12 เชี่ยงแซ (บท7)
    love_path = _find_docs("คู่สมพงษ์(ความรัก).xlsx")
    if love_path:
        love_wb = openpyxl.load_workbook(love_path, read_only=True, data_only=True)
        emit_qi_table(love_wb, "12เชี่ยงแซความรัก", "lover")
    return out


def main():
    os.makedirs(OUT, exist_ok=True)
    write("love-day-pillar.txt", extract_love_day_pillar())
    write("source7-custom.txt", extract_source7_custom())
    write("kheangkhung-reference.txt", extract_kheangkhung())
    write("career-relations.txt", extract_career_relations())


if __name__ == "__main__":
    main()
